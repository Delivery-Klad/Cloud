import os

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from fastapi import APIRouter, Cookie
from fastapi.responses import RedirectResponse
from typing import Optional
from pydantic import BaseModel
from git import Repo

from funcs.pages import show_forbidden_page
from funcs.utils import is_authorized_user

router = APIRouter(prefix="/add_text")
folder_path = "temp/files/7 сем/Информационно-поисковые системы/300 текстов"


class Test(BaseModel):
    main_theme: str
    arg: str
    themes: str
    link: str


@router.post("/")
async def add_text(query: Test, auth_psw: Optional[str] = Cookie(None)):
    try:
        if not is_authorized_user(auth_psw):
            return show_forbidden_page()
        document = Document()
        if len(query.arg.split(" ")) <= 200:
            path = f"{folder_path}/Short"
        elif 490 < len(query.arg.split(" ")) < 900:
            path = f"{folder_path}/Middle"
        elif 990 < len(query.arg.split(" ")) < 1600:
            path = f"{folder_path}/Long"
        texts = sorted(os.listdir(path), key=lambda x: int(x.split(".")[0]) if x.split(".")[0].isdigit() else 0)
        if len(texts) == 1:
            name = "1.docx"
        else:
            name = str(int(texts[len(texts) - 1].split(".")[0]) + 1) + ".docx"
        p = document.add_paragraph()
        font_styles = document.styles
        font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(14)
        font_object.name = 'Times New Roman'
        p.add_run(query.arg, style='CommentsStyle')
        document.save(f'{path}/{name}')

        document = Document()
        p = document.add_paragraph()
        font_styles = document.styles
        font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(14)
        font_object.name = 'Times New Roman'
        count = query.arg.replace("\n", " ")
        count = count.replace("- ", "")
        count = count.replace("  ", " ")
        count = count.replace("   ", " ")
        p.add_run(f"Количество слов - {len(count.split(' '))}\nОсновная тематика - {query.main_theme}\nСмежные"
                  f" тематики - {query.themes}\nИсточник - {query.link}", style='CommentsStyle')
        document.save(f'{path}/Справочные карточки/Справочная карточка_{name}')
        """repo = Repo("temp/.git")
        print("Untracked files detected...")
        repo.git.add(all=True)
        repo.index.commit("commit from cloud")
        origin = repo.remote(name='origin')
        origin.push()
        print("Push success!")"""
        return RedirectResponse(f"/files/7%20сем/Информационно-поисковые%20системы", status_code=302)
    except AttributeError:
        return show_forbidden_page()
    except Exception as er:
        print(er)


@router.get("/")
async def get_link(query: str):
    query = query.strip()
    res = []
    for i in os.listdir(f"{folder_path}/Short/Справочные карточки"):
        if i != "init":
            doc = Document(f"{folder_path}/Short/Справочные карточки/{i}")
            for p in doc.paragraphs:
                if str(p.text.split("\n")[3])[11:] == query:
                    res.append("Short/" + i)
    for i in os.listdir(f"{folder_path}/Middle/Справочные карточки"):
        if i != "init":
            doc = Document(f"{folder_path}/Middle/Справочные карточки/{i}")
            for p in doc.paragraphs:
                if str(p.text.split("\n")[3])[11:] == query:
                    res.append("Middle/" + i)
    for i in os.listdir(f"{folder_path}/Long/Справочные карточки"):
        if i != "init":
            doc = Document(f"{folder_path}/Long/Справочные карточки/{i}")
            for p in doc.paragraphs:
                if str(p.text.split("\n")[3])[11:] == query:
                    res.append("Long/" + i)
    if not res:
        return {"res": "Цыганства не обнаружено"}
    else:
        return {"res": str(res)[1:-1]}
